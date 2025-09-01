"""
Document Processing System for Inyandiko Legal AI Assistant
Handles all document types, formats, and edge cases with comprehensive processing capabilities.
Version: 3.2 (Enterprise Edition - Final Bug Fix Release)
"""

import os
import io
import re
import json
import hashlib
import logging
import asyncio
import tempfile
import base64
from typing import Dict, List, Optional, Tuple, Any, Union, Set
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass, field, asdict as dataclass_to_dict
from enum import Enum

# --- Core Document Processing Libraries ---
import fitz  # PyMuPDF
import pdfplumber
import PyPDF2
from PIL import Image, ImageEnhance, ImageFilter
import pytesseract
from docx import Document as DocxDocument
from docx.opc.exceptions import OpcError
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
import pandas as pd
from bs4 import BeautifulSoup
import markdown
import ebooklib
from ebooklib import epub
import zipfile
import rarfile
import py7zr


# --- Text Processing and NLP Libraries ---
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
import spacy
from langdetect import detect, DetectorFactory
import textstat
import textdistance
from textblob import TextBlob
import yake

# --- Machine Learning and Embeddings Libraries ---
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
from sklearn.decomposition import LatentDirichletAllocation
from sklearn.metrics.pairwise import cosine_similarity
import sentence_transformers

# --- Async, Performance, and Utility Libraries ---
import aiofiles
from concurrent.futures import ThreadPoolExecutor
import magic
import psutil
import time
from prometheus_client import Counter, Histogram, Gauge

# ==============================================================================
# 1. INITIAL SETUP
# ==============================================================================

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - [%(filename)s:%(lineno)d] - %(message)s",
)
logger = logging.getLogger(__name__)

# Set deterministic language detection for reproducibility
DetectorFactory.seed = 0


def setup_nltk():
    """Downloads required NLTK data if not present."""
    required_data = ["punkt", "stopwords", "averaged_perceptron_tagger"]
    for package in required_data:
        try:
            # Correctly check for packages without the 'quiet' parameter
            if package in ("punkt", "stopwords"):
                nltk.data.find(f"tokenizers/{package}")
            else:
                nltk.data.find(f"taggers/{package}")
        except LookupError:
            logging.info(f"Downloading NLTK package: {package}")
            nltk.download(package)


setup_nltk()

# Initialize spaCy model for NER and other NLP tasks
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    logging.warning(
        "spaCy 'en_core_web_sm' model not found. NER and some NLP features will be disabled."
    )
    nlp = None

# ==============================================================================
# 2. METRICS DEFINITION (for Prometheus)
# ==============================================================================

doc_processing_counter = Counter(
    "docproc_documents_processed_total", "Total documents processed", ["status", "type"]
)
doc_processing_duration = Histogram(
    "docproc_processing_duration_seconds", "Document processing duration"
)
doc_size_gauge = Gauge("docproc_document_size_bytes", "Document size in bytes")
cpu_usage_gauge = Gauge("docproc_cpu_percent", "CPU usage during document processing")
memory_usage_gauge = Gauge(
    "docproc_memory_mb", "Memory usage in MB during document processing"
)

# ==============================================================================
# 3. CONFIGURATION, ENUMS, AND DATA CLASSES
# ==============================================================================


@dataclass
class DocumentProcessorConfig:
    """Centralized configuration for the document processor."""

    max_file_size_mb: int = 100
    max_pages_per_doc: int = 1000
    processing_timeout_seconds: int = 300
    cache_size_limit: int = 100
    ocr_confidence_threshold: int = 60
    default_chunk_size: int = 1000
    default_chunk_overlap: int = 200
    semantic_similarity_threshold: float = 0.75
    num_topics_for_chunking: int = 10
    max_summary_sentences: int = 5
    max_keywords: int = 10
    pii_redaction_enabled: bool = True


class DocumentType(Enum):
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    HTML = "html"
    MARKDOWN = "md"
    XLSX = "xlsx"
    CSV = "csv"
    JSON = "json"
    XML = "xml"
    EPUB = "epub"
    ZIP = "zip"
    RAR = "rar"
    SEVEN_Z = "7z"
    IMAGE = "image"
    UNKNOWN = "unknown"


class ProcessingStatus(Enum):
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"
    CORRUPTED = "corrupted"
    UNSUPPORTED = "unsupported"
    SECURITY_RISK = "security_risk"


class ChunkingStrategy(Enum):
    FIXED_SIZE = "fixed_size"
    SENTENCE_BASED = "sentence_based"
    PARAGRAPH_BASED = "paragraph_based"
    SEMANTIC_BASED = "semantic_based"
    HIERARCHICAL = "hierarchical"
    TOPIC_BASED = "topic_based"


@dataclass
class DocumentMetadata:
    """Comprehensive, structured metadata for a processed document."""

    file_path: str
    file_name: str
    file_size: int
    file_hash: str
    mime_type: str
    document_type: DocumentType
    created_at: datetime
    modified_at: datetime
    processed_at: Optional[datetime] = None
    language: Optional[str] = None
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    character_count: Optional[int] = None
    reading_level: Optional[float] = None
    sentiment: Optional[Dict[str, float]] = None
    summary: Optional[str] = None
    keywords: List[str] = field(default_factory=list)
    topics: List[str] = field(default_factory=list)
    entities: List[Dict] = field(default_factory=list)
    pii_detected: List[Dict] = field(default_factory=list)
    security_scan_result: Optional[Dict] = None
    processing_time: Optional[float] = None
    resource_usage: Optional[Dict[str, float]] = None
    error_message: Optional[str] = None
    status: ProcessingStatus = ProcessingStatus.PENDING


@dataclass
class DocumentChunk:
    """A structured representation of a single chunk of document content."""

    chunk_id: str
    content: str
    chunk_index: int
    start_position: int
    end_position: int
    page_number: Optional[int] = None
    section_title: Optional[str] = None
    chunk_type: str = "text"
    confidence_score: float = 1.0
    metadata: Dict[str, Any] = field(default_factory=dict)
    embedding: Optional[np.ndarray] = None


@dataclass
class ProcessedDocument:
    """The final, complete output of the document processing pipeline."""

    metadata: DocumentMetadata
    chunks: List[DocumentChunk]
    full_text: str
    redacted_text: Optional[str] = None
    structure: Optional[Dict] = None
    images: List[Dict] = field(default_factory=list)
    tables: List[Dict] = field(default_factory=list)
    links: List[str] = field(default_factory=list)
    annotations: List[str] = field(default_factory=list)


# ==============================================================================
# 4. CORE COMPONENT CLASSES
# ==============================================================================


class SecurityScanner:
    """Scans documents for potential security risks like macros, scripts, and large sizes."""

    def __init__(self, config: DocumentProcessorConfig):
        self.config = config
        self.max_file_size = self.config.max_file_size_mb * 1024 * 1024
        self.blocked_extensions = {
            ".exe",
            ".bat",
            ".cmd",
            ".scr",
            ".pif",
            ".com",
            ".vbs",
        }
        self.suspicious_patterns = [
            re.compile(p, re.IGNORECASE)
            for p in [
                r"javascript:",
                r"vbscript:",
                r"<script",
                r"eval\(",
                r"document\.write",
                r"window\.location",
            ]
        ]

    async def scan_file(self, file_path: str) -> Dict[str, Any]:
        scan_result: Dict[str, Any] = {"safe": True, "threats": [], "warnings": []}
        try:
            file_size = os.path.getsize(file_path)
            if file_size > self.max_file_size:
                scan_result.update(
                    {
                        "safe": False,
                        "threats": [
                            f"File exceeds size limit of {self.config.max_file_size_mb}MB"
                        ],
                    }
                )
                return scan_result

            if Path(file_path).suffix.lower() in self.blocked_extensions:
                scan_result.update(
                    {"safe": False, "threats": ["File has a blocked extension."]}
                )
                return scan_result

            scan_result["mime_type"] = magic.from_file(file_path, mime=True)
            try:
                async with aiofiles.open(
                    file_path, "r", encoding="utf-8", errors="ignore"
                ) as f:
                    content_sample = await f.read(20480)  # Scan first 20KB
                    for pattern in self.suspicious_patterns:
                        if pattern.search(content_sample):
                            scan_result["warnings"].append(
                                f"Suspicious pattern found: {pattern.pattern}"
                            )
            except (IOError, UnicodeDecodeError):
                pass  # Ignore errors for binary files
        except Exception as e:
            scan_result["warnings"].append(f"Security scan error: {e}")
        return scan_result


class OCRProcessor:
    """Handles OCR for images and scanned documents using Tesseract."""

    def __init__(self, config: DocumentProcessorConfig):
        self.config = config

    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        if image.mode != "L":
            image = image.convert("L")
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(2.0)
        image = image.filter(ImageFilter.SHARPEN)
        return image

    async def extract_text_from_image_async(
        self, image_bytes: bytes, lang: str = "eng"
    ) -> Dict[str, Any]:
        loop = asyncio.get_running_loop()

        def ocr_task():
            try:
                image = Image.open(io.BytesIO(image_bytes))
                processed_image = self._preprocess_image(image)
                data = pytesseract.image_to_data(
                    processed_image, lang=lang, output_type=pytesseract.Output.DICT
                )
                text_parts, confidences = [], []
                for i, conf_str in enumerate(data["conf"]):
                    conf = int(float(conf_str)) if conf_str != "-1" else -1
                    if conf > self.config.ocr_confidence_threshold:
                        text = data["text"][i].strip()
                        if text:
                            text_parts.append(text)
                            confidences.append(conf)
                return {
                    "text": " ".join(text_parts),
                    "confidence": np.mean(confidences) if confidences else 0.0,
                    "success": True,
                }
            except Exception as e:
                logging.error(f"OCR task failed: {e}")
                return {
                    "text": "",
                    "confidence": 0.0,
                    "success": False,
                    "error": str(e),
                }

        return await loop.run_in_executor(None, ocr_task)


class TextAnalyzer:
    """Performs comprehensive linguistic analysis of text."""

    def __init__(self, config: DocumentProcessorConfig):
        self.config = config
        self.stop_words = set(stopwords.words("english"))
        self.sentence_model: Optional[sentence_transformers.SentenceTransformer] = None
        self.kw_extractor = yake.KeywordExtractor(
            lan="en", n=3, dedupLim=0.9, top=self.config.max_keywords, features=None
        )

        self.pii_patterns = {
            "EMAIL": re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"),
            "PHONE": re.compile(r"\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}"),
        }

    async def initialize(self):
        """Initializes heavy components like ML models asynchronously."""
        if self.sentence_model is not None:
            return

        logger.info("Initializing TextAnalyzer: loading SentenceTransformer model...")
        try:
            loop = asyncio.get_running_loop()
            # Use a default executor for the blocking model load
            self.sentence_model = await loop.run_in_executor(
                None,
                sentence_transformers.SentenceTransformer,
                "all-MiniLM-L6-v2",
            )
            logger.info("TextAnalyzer initialized successfully.")
        except Exception as e:
            logging.warning(
                f"Could not load SentenceTransformer model: {e}. Semantic features will be disabled."
            )
            self.sentence_model = None

    def analyze(self, text: str) -> Dict[str, Any]:
        """Runs a full suite of text analyses."""
        if not text:
            return {}
        return {
            "language": self.detect_language(text),
            "reading_level": self.calculate_reading_level(text),
            "word_count": len(word_tokenize(text)),
            "character_count": len(text),
            "sentiment": self.analyze_sentiment(text),
            "summary": self.summarize_text(text),
            "keywords": self.extract_keywords(text),
            "entities": self.extract_entities(text),
            "pii": self.detect_pii(text),
        }

    def detect_language(self, text: str) -> str:
        try:
            return detect(text) if len(text.strip()) > 20 else "unknown"
        except Exception:
            return "unknown"

    def calculate_reading_level(self, text: str) -> float:
        try:
            return textstat.textstat.flesch_reading_ease(text)
        except Exception:
            return 0.0

    def analyze_sentiment(self, text: str) -> Dict[str, float]:
        try:
            blob = TextBlob(text)
            sentiment_result = blob.sentiment
            return {"polarity": sentiment_result.polarity, "subjectivity": sentiment_result.subjectivity}  # type: ignore
        except Exception:
            return {"polarity": 0.0, "subjectivity": 0.0}

    def summarize_text(self, text: str) -> str:
        try:
            sentences = sent_tokenize(text)
            if len(sentences) <= self.config.max_summary_sentences:
                return " ".join(sentences)
            vectorizer = TfidfVectorizer(stop_words="english")
            tfidf_matrix = vectorizer.fit_transform(sentences)
            sentence_scores = tfidf_matrix.toarray().sum(axis=1).flatten()  # type: ignore
            top_indices = sentence_scores.argsort()[
                -self.config.max_summary_sentences :
            ][::-1]
            summary_sentences = [sentences[i] for i in sorted(top_indices)]
            return " ".join(summary_sentences)
        except Exception:
            return text[:500]

    def extract_keywords(self, text: str) -> List[str]:
        try:
            return [kw for kw, score in self.kw_extractor.extract_keywords(text)]
        except Exception:
            return []

    def extract_entities(self, text: str) -> List[Dict[str, Any]]:
        if not nlp:
            return []
        try:
            doc = nlp(text[:1_000_000])
            return [{"text": ent.text, "label": ent.label_} for ent in doc.ents]
        except Exception:
            return []

    def detect_pii(self, text: str) -> List[Dict[str, Any]]:
        pii_found = []
        for pii_type, pattern in self.pii_patterns.items():
            for match in pattern.finditer(text):
                pii_found.append(
                    {
                        "text": match.group(0),
                        "type": pii_type,
                        "start": match.start(),
                        "end": match.end(),
                    }
                )
        return pii_found

    def redact_text(self, text: str, pii_list: List[Dict[str, Any]]) -> str:
        """Redacts detected PII from text."""
        if not self.config.pii_redaction_enabled or not pii_list:
            return text

        pii_list.sort(key=lambda p: p["end"], reverse=True)

        for pii in pii_list:
            text = text[: pii["start"]] + f"[{pii['type']}]" + text[pii["end"] :]
        return text

    def calculate_similarity(self, text1: str, text2: str) -> float:
        if not self.sentence_model:
            return float(
                textdistance.jaccard.similarity(
                    word_tokenize(text1), word_tokenize(text2)
                )
            )
        try:
            embeddings = self.sentence_model.encode([text1, text2])
            return float(
                cosine_similarity(
                    np.array(embeddings[0]).reshape(1, -1),
                    np.array(embeddings[1]).reshape(1, -1),
                )[0][0]
            )
        except Exception:
            return 0.0


class ChunkingEngine:
    """Provides multiple advanced strategies for splitting documents into chunks."""

    def __init__(self, text_analyzer: TextAnalyzer, config: DocumentProcessorConfig):
        self.text_analyzer = text_analyzer
        self.config = config

    def chunk(
        self, strategy: ChunkingStrategy, text: str, structure: Optional[Dict] = None
    ) -> List[DocumentChunk]:
        """Selects and applies a chunking strategy."""
        if strategy == ChunkingStrategy.TOPIC_BASED:
            return self._chunk_by_topic(text)
        if (
            strategy == ChunkingStrategy.HIERARCHICAL
            and structure
            and structure.get("toc")
        ):
            return self._chunk_hierarchically(text, structure)

        chunks = self._chunk_by_paragraphs(text)
        if not chunks:
            chunks = self._chunk_by_fixed_size(text)
        return chunks

    def _chunk_by_paragraphs(self, text: str) -> List[DocumentChunk]:
        paragraphs = [p.strip() for p in text.split("\n\n") if len(p.strip()) > 50]
        chunks = []
        for i, p in enumerate(paragraphs):
            try:
                start_pos = text.index(p)
                chunks.append(
                    DocumentChunk(
                        chunk_id=f"para_{i}",
                        content=p,
                        chunk_index=i,
                        start_position=start_pos,
                        end_position=start_pos + len(p),
                    )
                )
            except ValueError:
                continue
        return chunks

    def _chunk_by_fixed_size(self, text: str) -> List[DocumentChunk]:
        size = self.config.default_chunk_size
        overlap = self.config.default_chunk_overlap
        chunks, start, index = [], 0, 0
        while start < len(text):
            end = min(start + size, len(text))
            content = text[start:end].strip()
            if content:
                chunks.append(
                    DocumentChunk(
                        chunk_id=f"fixed_{index}",
                        content=content,
                        chunk_index=index,
                        start_position=start,
                        end_position=end,
                    )
                )
                index += 1
            start += size - overlap
        return chunks

    def _chunk_hierarchically(self, text: str, structure: Dict) -> List[DocumentChunk]:
        chunks = []
        toc = structure.get("toc", [])
        if not toc:
            return self._chunk_by_paragraphs(text)

        content_map = []
        for i, (level, title, page_num) in enumerate(toc):
            start_pos = text.find(title)
            if start_pos != -1:
                content_map.append(
                    {"title": title, "start": start_pos, "page": page_num}
                )

        content_map.sort(key=lambda x: x["start"])

        for i, section in enumerate(content_map):
            start_pos = section["start"]
            end_pos = (
                content_map[i + 1]["start"] if i + 1 < len(content_map) else len(text)
            )
            content = text[start_pos:end_pos].strip()

            if len(content) > 50:
                chunks.append(
                    DocumentChunk(
                        chunk_id=f"hier_{i}",
                        content=content,
                        chunk_index=len(chunks),
                        start_position=start_pos,
                        end_position=end_pos,
                        section_title=section["title"],
                        page_number=section["page"],
                    )
                )
        return chunks

    def _chunk_by_topic(self, text: str) -> List[DocumentChunk]:
        paragraphs = [p for p in text.split("\n\n") if len(p.strip()) > 100]
        if len(paragraphs) < self.config.num_topics_for_chunking:
            return self._chunk_by_paragraphs(text)
        try:
            vectorizer = CountVectorizer(
                stop_words="english", max_features=1000, min_df=2
            )
            X = vectorizer.fit_transform(paragraphs)
            lda = LatentDirichletAllocation(
                n_components=self.config.num_topics_for_chunking, random_state=42
            )
            topic_assignments = lda.fit_transform(X).argmax(axis=1)

            chunks, current_chunk_paras, current_topic = [], [], topic_assignments[0]
            for i, para in enumerate(paragraphs):
                if (
                    topic_assignments[i] == current_topic
                    and len("\n\n".join(current_chunk_paras)) < 4000
                ):
                    current_chunk_paras.append(para)
                else:
                    content = "\n\n".join(current_chunk_paras)
                    chunks.append(
                        DocumentChunk(
                            chunk_id=f"topic_{len(chunks)}",
                            content=content,
                            chunk_index=len(chunks),
                            start_position=text.find(content),
                            end_position=text.find(content) + len(content),
                            metadata={"topic_id": int(current_topic)},
                        )
                    )
                    current_chunk_paras, current_topic = [para], topic_assignments[i]

            if current_chunk_paras:
                content = "\n\n".join(current_chunk_paras)
                chunks.append(
                    DocumentChunk(
                        chunk_id=f"topic_{len(chunks)}",
                        content=content,
                        chunk_index=len(chunks),
                        start_position=text.find(content),
                        end_position=text.find(content) + len(content),
                        metadata={"topic_id": int(current_topic)},
                    )
                )
            return chunks
        except Exception as e:
            logging.warning(
                f"Topic-based chunking failed: {e}. Falling back to paragraphs."
            )
            return self._chunk_by_paragraphs(text)


# ==============================================================================
# 5. DOCUMENT EXTRACTOR CLASSES
# ==============================================================================


class BaseExtractor:
    """Abstract base class for all document-type-specific extractors."""

    def __init__(
        self,
        config: DocumentProcessorConfig,
        ocr_processor: Optional[OCRProcessor] = None,
    ):
        self.config = config
        self.ocr_processor = ocr_processor

    async def extract(self, file_path: str) -> Tuple[str, Dict, List, List, List, List]:
        raise NotImplementedError


class PdfExtractor(BaseExtractor):
    async def extract(self, file_path: str) -> Tuple[str, Dict, List, List, List, List]:
        text = ""
        structure: Dict[str, Any] = {"toc": [], "links": [], "annotations": []}
        images, tables, links, annotations = [], [], [], []

        with fitz.Document(file_path) as doc:
            structure["toc"] = doc.get_toc()  # type: ignore
            structure["page_count"] = len(doc)
            text_builder = []
            for i in range(len(doc)):
                if i >= self.config.max_pages_per_doc:
                    break
                page = doc.load_page(i)
                text_builder.append(page.get_text())
                links.extend(
                    [link["uri"] for link in page.get_links() if "uri" in link]
                )
                annotations.extend(
                    [
                        annot.info["content"]
                        for annot in page.annots()
                        if annot.info.get("content")
                    ]
                )

                for img_instance in page.get_images(full=True):
                    xref = img_instance[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    images.append(
                        {
                            "page": i + 1,
                            "bytes": image_bytes,
                            "ext": base_image["ext"],
                            "b64": base64.b64encode(image_bytes).decode("utf-8"),
                        }
                    )
            text = "\n\n".join(text_builder)

        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                if i >= self.config.max_pages_per_doc:
                    break
                page_tables = page.extract_tables()
                if page_tables:
                    for tbl in page_tables:
                        tables.append(
                            {
                                "page": i + 1,
                                "data_html": pd.DataFrame(
                                    tbl[1:], columns=tbl[0]
                                ).to_html(),
                            }
                        )

        return text.strip(), structure, images, tables, links, annotations


class DocxExtractor(BaseExtractor):
    async def extract(self, file_path: str) -> Tuple[str, Dict, List, List, List, List]:
        try:
            doc = DocxDocument(file_path)
            text_builder, tables = [], []
            for para in doc.paragraphs:
                text_builder.append(para.text)
            for table in doc.tables:
                table_data = [[cell.text for cell in row.cells] for row in table.rows]
                tables.append(
                    {"page": None, "data_html": pd.DataFrame(table_data).to_html()}
                )
                text_builder.append(pd.DataFrame(table_data).to_string())

            props = doc.core_properties
            structure = {
                "author": props.author,
                "title": props.title,
                "created": props.created,
            }
            return "\n".join(text_builder).strip(), structure, [], tables, [], []
        except OpcError as e:
            raise IOError(f"File {file_path} is not a valid DOCX file.") from e


class SpreadsheetExtractor(BaseExtractor):
    async def extract(self, file_path: str) -> Tuple[str, Dict, List, List, List, List]:
        xls = pd.ExcelFile(file_path)
        text_builder, tables = [], []
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            if not df.empty:
                text_builder.append(
                    f"--- SHEET: {sheet_name} ---\n{df.to_string()}\n\n"
                )
                tables.append({"page": sheet_name, "data_html": df.to_html()})
        return (
            "\n".join(text_builder).strip(),
            {"sheets": xls.sheet_names},
            [],
            tables,
            [],
            [],
        )


class TextualExtractor(BaseExtractor):
    """Handles plain text, HTML, Markdown, etc."""

    async def extract(self, file_path: str) -> Tuple[str, Dict, List, List, List, List]:
        async with aiofiles.open(
            file_path, "r", encoding="utf-8", errors="ignore"
        ) as f:
            content = await f.read()

        ext = Path(file_path).suffix.lower()
        if ext in [".html", ".htm"]:
            soup = BeautifulSoup(content, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            text = soup.get_text(separator="\n", strip=True)
        elif ext == ".md":
            html = markdown.markdown(content)
            text = BeautifulSoup(html, "html.parser").get_text(
                separator="\n", strip=True
            )
        else:
            text = content

        return text, {}, [], [], [], []


class ArchiveExtractor(BaseExtractor):
    async def extract(self, file_path: str) -> Tuple[str, Dict, List, List, List, List]:
        text_builder, structure = [], {"files": []}

        with tempfile.TemporaryDirectory() as temp_dir:
            try:
                if file_path.endswith(".zip"):
                    with zipfile.ZipFile(file_path, "r") as zf:
                        zf.extractall(temp_dir)
                elif file_path.endswith(".rar"):
                    with rarfile.RarFile(file_path, "r") as rf:
                        rf.extractall(temp_dir)
                elif file_path.endswith(".7z"):
                    with py7zr.SevenZipFile(file_path, mode="r") as z:
                        z.extractall(temp_dir)
            except Exception as e:
                raise IOError(f"Failed to extract archive {file_path}") from e

            for root, _, files in os.walk(temp_dir):
                for file in files:
                    structure["files"].append(file)
            text_builder.append(
                f"Archive contains files: {', '.join(structure['files'])}"
            )

        return "\n".join(text_builder), structure, [], [], [], []


# ==============================================================================
# 6. MAIN PROCESSOR CLASS
# ==============================================================================


class DocumentProcessor:
    """Orchestrates the end-to-end document processing pipeline."""

    def __init__(self, config: Optional[DocumentProcessorConfig] = None):
        self.config = config or DocumentProcessorConfig()
        self.ocr_processor = OCRProcessor(self.config)
        self.security_scanner = SecurityScanner(self.config)
        self.text_analyzer = TextAnalyzer(self.config)
        self.chunking_engine = ChunkingEngine(self.text_analyzer, self.config)
        self.executor = ThreadPoolExecutor(max_workers=os.cpu_count() or 4)
        self.document_cache: Dict[str, ProcessedDocument] = {}
        self.is_initialized = False

        self.extractor_map = {
            DocumentType.PDF: PdfExtractor(self.config),
            DocumentType.DOCX: DocxExtractor(self.config),
            DocumentType.XLSX: SpreadsheetExtractor(self.config),
            DocumentType.TXT: TextualExtractor(self.config),
            DocumentType.HTML: TextualExtractor(self.config),
            DocumentType.MARKDOWN: TextualExtractor(self.config),
            DocumentType.ZIP: ArchiveExtractor(self.config),
            DocumentType.RAR: ArchiveExtractor(self.config),
            DocumentType.SEVEN_Z: ArchiveExtractor(self.config),
        }

    async def initialize(self):
        """
        Initializes the document processor and its sub-components, particularly
        those that involve loading heavy resources like ML models.
        """
        if self.is_initialized:
            logger.info("DocumentProcessor is already initialized.")
            return

        logger.info("Initializing DocumentProcessor...")
        # The most important sub-component to initialize is the TextAnalyzer,
        # as it loads a sentence-transformer model.
        await self.text_analyzer.initialize()

        # We can also perform a quick health check of dependencies here.
        try:
            pytesseract.get_tesseract_version()
        except pytesseract.TesseractNotFoundError:
            logger.error(
                "Tesseract is not installed or not in your PATH. OCR will fail."
            )

        if nlp is None:
            logger.warning(
                "spaCy model 'en_core_web_sm' not loaded. NER features will be disabled."
            )

        self.is_initialized = True
        logger.info("DocumentProcessor initialized successfully.")

    def _get_extractor(self, doc_type: DocumentType) -> Optional[BaseExtractor]:
        return self.extractor_map.get(doc_type)

    def _calculate_file_hash(self, file_path: str) -> str:
        sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            while chunk := f.read(8192):
                sha256.update(chunk)
        return sha256.hexdigest()

    def _detect_doc_type(self, file_path: str) -> Tuple[DocumentType, str]:
        mime_type_map = {
            "application/pdf": DocumentType.PDF,
            "text/plain": DocumentType.TXT,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DocumentType.DOCX,
            "text/html": DocumentType.HTML,
            "text/markdown": DocumentType.MARKDOWN,
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": DocumentType.XLSX,
            "application/zip": DocumentType.ZIP,
            "application/x-rar-compressed": DocumentType.RAR,
            "application/x-7z-compressed": DocumentType.SEVEN_Z,
        }
        try:
            mime_type = magic.from_file(file_path, mime=True)
            return mime_type_map.get(mime_type, DocumentType.UNKNOWN), mime_type
        except Exception:
            return DocumentType.UNKNOWN, "application/octet-stream"

    async def _create_metadata(self, file_path: str) -> DocumentMetadata:
        stat = os.stat(file_path)
        doc_type, mime = self._detect_doc_type(file_path)
        metadata = DocumentMetadata(
            file_path=file_path,
            file_name=Path(file_path).name,
            file_size=stat.st_size,
            file_hash=self._calculate_file_hash(file_path),
            mime_type=mime,
            document_type=doc_type,
            created_at=datetime.fromtimestamp(stat.st_ctime),
            modified_at=datetime.fromtimestamp(stat.st_mtime),
        )
        metadata.security_scan_result = await self.security_scanner.scan_file(file_path)
        if not metadata.security_scan_result["safe"]:
            metadata.status = ProcessingStatus.SECURITY_RISK
        return metadata

    async def process_document(
        self,
        file_path: str,
        chunking_strategy: ChunkingStrategy = ChunkingStrategy.HIERARCHICAL,
    ) -> ProcessedDocument:
        """Main processing pipeline for a single document."""
        process = psutil.Process(os.getpid())
        start_mem = process.memory_info().rss / (1024 * 1024)
        start_time = time.monotonic()

        metadata = DocumentMetadata(
            file_path,
            Path(file_path).name,
            0,
            "",
            "application/octet-stream",
            DocumentType.UNKNOWN,
            datetime.now(),
            datetime.now(),
        )
        text = ""
        structure: Optional[Dict] = None
        images: List[Dict] = []
        tables: List[Dict] = []
        links: List[str] = []
        annotations: List[str] = []
        redacted_text: Optional[str] = None
        chunks: List[DocumentChunk] = []
        cache_key: Optional[str] = None

        if not self.is_initialized:
            logger.error(
                "DocumentProcessor has not been initialized. Call `await processor.initialize()` before use."
            )
            raise RuntimeError(
                "DocumentProcessor has not been initialized. Call `await processor.initialize()` before use."
            )

        try:
            metadata = await self._create_metadata(file_path)
            if metadata.status == ProcessingStatus.SECURITY_RISK:
                return ProcessedDocument(metadata, [], "")

            cache_key = f"{metadata.file_hash}_{chunking_strategy.value}"
            if cache_key in self.document_cache:
                return self.document_cache[cache_key]

            extractor = self._get_extractor(metadata.document_type)
            if not extractor:
                raise NotImplementedError(
                    f"No extractor for type {metadata.document_type.value}"
                )

            metadata.status = ProcessingStatus.PROCESSING
            text, structure, images, tables, links, annotations = (
                await extractor.extract(file_path)
            )

            if not text.strip():
                metadata.status = ProcessingStatus.FAILED
                metadata.error_message = "No text content extracted"
                return ProcessedDocument(
                    metadata,
                    [],
                    text,
                    structure=structure,
                    images=images,
                    tables=tables,
                    links=links,
                    annotations=annotations,
                )

            loop = asyncio.get_running_loop()
            text_analysis = await loop.run_in_executor(
                self.executor, self.text_analyzer.analyze, text
            )

            for key, value in text_analysis.items():
                if hasattr(metadata, key):
                    setattr(metadata, key, value)
            metadata.pii_detected = text_analysis.get("pii", [])
            redacted_text = self.text_analyzer.redact_text(text, metadata.pii_detected)

            chunks = self.chunking_engine.chunk(chunking_strategy, text, structure)
            metadata.status = ProcessingStatus.COMPLETED

        except Exception as e:
            logging.error(f"Failed to process {file_path}: {e}", exc_info=True)
            metadata.status = ProcessingStatus.FAILED
            metadata.error_message = str(e)
        finally:
            end_mem = process.memory_info().rss / (1024 * 1024)
            metadata.processing_time = time.monotonic() - start_time
            metadata.resource_usage = {"mem_usage_mb": end_mem - start_mem}
            doc_processing_counter.labels(
                status=metadata.status.value, type=metadata.document_type.value
            ).inc()

        processed_doc = ProcessedDocument(
            metadata=metadata,
            chunks=chunks,
            full_text=text,
            redacted_text=redacted_text,
            structure=structure,
            images=images,
            tables=tables,
            links=links,
            annotations=annotations,
        )

        if cache_key is not None:
            if len(self.document_cache) >= self.config.cache_size_limit:
                self.document_cache.pop(next(iter(self.document_cache)))
            self.document_cache[cache_key] = processed_doc

        return processed_doc

    async def batch_process(self, file_paths: List[str]) -> List[ProcessedDocument]:
        """Processes a batch of documents concurrently."""
        tasks = [self.process_document(fp) for fp in file_paths]
        return await asyncio.gather(*tasks)

    def health_check(self) -> bool:
        """Checks the health of the DocumentProcessor and its dependencies."""
        healthy = True
        logger.info("Running DocumentProcessor health check...")
        loop = asyncio.get_running_loop()

        # 1. Check Tesseract OCR
        try:
            version = pytesseract.get_tesseract_version()
            logger.info(f"Tesseract version {version} found.")
        except pytesseract.TesseractNotFoundError:
            logger.error("Tesseract is not installed or not in your PATH.")
            healthy = False
        except Exception as e:
            logger.error(f"Error checking Tesseract: {e}", exc_info=True)
            healthy = False

        # 2. Check spaCy model
        if nlp is None:
            logger.error("spaCy 'en_core_web_sm' model is not loaded.")
            healthy = False
        else:
            logger.info("spaCy 'en_core_web_sm' model is loaded.")

        # 3. Check magic library
        try:
            # This is a fast, local operation, so running it directly is acceptable for a health check.
            with tempfile.NamedTemporaryFile(mode='w', delete=True, suffix=".txt") as tmp:
                tmp.write("test")
                tmp.flush()
                mime_type = magic.from_file(tmp.name, mime=True)
                if mime_type == "text/plain":
                    logger.info("Magic library is functioning correctly.")
                else:
                    logger.warning(f"Magic library returned unexpected mime type for .txt: {mime_type}")
        except Exception as e:
            logger.error(f"Magic library check failed: {e}", exc_info=True)
            healthy = False

        # 4. Check ThreadPoolExecutor
        if self.executor._shutdown:
            logger.error("ThreadPoolExecutor is shut down.")
            healthy = False
        else:
            logger.info("ThreadPoolExecutor is active.")

        if healthy:
            logger.info("DocumentProcessor health check passed.")
        else:
            logger.error("DocumentProcessor health check failed.")
            
        return healthy

    async def cleanup(self):
        self.executor.shutdown()
        logging.info("Document processor resources have been cleaned up.")


# ==============================================================================
# 7. OUTPUT & REPORTING UTILITY
# ==============================================================================


class ResultExporter:
    """Exports processed document data to structured formats like JSON."""

    @staticmethod
    async def to_json(doc: ProcessedDocument, output_path: str):
        class CustomEncoder(json.JSONEncoder):
            def default(self, o):
                if isinstance(o, (datetime,)):
                    return o.isoformat()
                if isinstance(o, Enum):
                    return o.value
                if isinstance(o, np.ndarray):
                    return o.tolist()
                if isinstance(o, bytes):
                    return "<binary_data>"
                return super().default(o)

        data = {
            "metadata": dataclass_to_dict(doc.metadata),
            "full_text_snippet": doc.full_text[:500] + "...",
            "structure": doc.structure,
            "tables_count": len(doc.tables),
            "images_count": len(doc.images),
            "chunks": [dataclass_to_dict(c) for c in doc.chunks],
        }

        async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
            await f.write(json.dumps(data, indent=2, cls=CustomEncoder))


def print_processing_report(doc: ProcessedDocument):
    """Prints a rich, formatted report of a processed document using the 'rich' library."""
    try:
        from rich.console import Console
        from rich.table import Table
        from rich.panel import Panel
        from rich.syntax import Syntax

        console = Console()
        status_color = {
            "COMPLETED": "green",
            "FAILED": "red",
            "SECURITY_RISK": "yellow",
        }.get(doc.metadata.status.value, "white")

        title = f"[{status_color}]Processing Report for: [bold cyan]{doc.metadata.file_name}[/bold cyan]"

        meta_table = Table(show_header=False, box=None, padding=(0, 1))
        meta_table.add_column(style="magenta", justify="right")
        meta_table.add_column(style="green")
        meta_table.add_row(
            "Status:", f"[{status_color}]{doc.metadata.status.value}[/{status_color}]"
        )
        if doc.metadata.error_message:
            meta_table.add_row("Error:", f"[red]{doc.metadata.error_message}")
        meta_table.add_row(
            "Type:", f"{doc.metadata.document_type.value} ({doc.metadata.mime_type})"
        )
        meta_table.add_row("Language:", str(doc.metadata.language))
        meta_table.add_row(
            "Counts:",
            f"{doc.metadata.word_count} words, {doc.metadata.character_count} chars",
        )
        if doc.metadata.processing_time is not None:
            meta_table.add_row(
                "Processing Time:", f"{doc.metadata.processing_time:.3f}s"
            )
        if doc.metadata.resource_usage:
            meta_table.add_row(
                "Memory Usage:", f"{doc.metadata.resource_usage['mem_usage_mb']:.2f} MB"
            )

        analysis_text = (
            f"[bold]Summary:[/bold]\n{doc.metadata.summary}\n\n"
            f"[bold]Keywords:[/bold] {', '.join(doc.metadata.keywords)}\n"
            f"[bold]Sentiment:[/bold] Polarity={(doc.metadata.sentiment or {}).get('polarity', 0):.2f}, Subjectivity={(doc.metadata.sentiment or {}).get('subjectivity', 0):.2f}"
            f"[bold]\nPII Detected:[/bold] {[pii['type'] for pii in doc.metadata.pii_detected]}"
        )

        chunks_text = f"Total Chunks: {len(doc.chunks)}\n"
        if doc.chunks:
            chunks_text += f"\n[bold]First Chunk Preview (ID: {doc.chunks[0].chunk_id}):[/bold]\n'{doc.chunks[0].content[:200]}...'"

        console.print(Panel(meta_table, title=title, border_style="blue", expand=False))
        if doc.metadata.status == ProcessingStatus.COMPLETED:
            console.print(
                Panel(
                    analysis_text,
                    title="[green]Text Analysis[/green]",
                    border_style="green",
                    expand=False,
                )
            )
            console.print(
                Panel(
                    chunks_text,
                    title="[yellow]Chunking[/yellow]",
                    border_style="yellow",
                    expand=False,
                )
            )

        console.print("-" * console.width)

    except ImportError:
        logging.warning("Install 'rich' (`pip install rich`) for formatted reports.")
        print(f"\n--- Report for {doc.metadata.file_name} ---")
        print(f"Status: {doc.metadata.status.value}")
        if doc.metadata.error_message:
            print(f"Error: {doc.metadata.error_message}")
